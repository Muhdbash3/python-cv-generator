from docx import Document
import pyttsx3


def speak(text):
    pyttsx3.speak(text)


d = Document()
name = input('what is your name? ')
speak("hello" + name + 'how is your day?')
speak('Type your phone number')
phone_number = input('Type your phone number ')
email = input('what is your email? ')

d.add_paragraph(name + ' ' + '|' + phone_number + ' ' + "|" + email)

d.add_heading('About Me')
d.add_paragraph(input('Tell us more  about your self '))

d.add_heading('Work Experiences')
print('Tell us your work experience')
company = input('what is the name of your company? ')
post_held = input('The post held ')
Start_date = input("Start date ")
End_date = input('Stop date ')

p = d.add_paragraph()

p.add_run(company + ' ').bold = True
p.add_run(Start_date + '-' + End_date + '\n').italic = True
p.add_run(post_held)

while True:
    has_other_experience = input('Do you have another work experience? Yes or No ')
    if has_other_experience.lower() == 'yes':
        print('Tell us your work experience')
        company = input('what is your company name? ')
        post_held = input('The post held ')
        Start_date = input("Start date... ")
        End_date = input('Stop date... ')

        p = d.add_paragraph()
        p.add_run(company + ' ').bold = True
        p.add_run(Start_date + '-' + End_date + '\n').italic = True
        p.add_run(post_held)
    else:
        break

d.add_heading('Skill(s)')
speak('Add your skills')
skills = input('Add your skills ')
p = d.add_paragraph(skills)
p.style = 'List Bullet'

while True:
    has_another_skill = input('Do you have another Skill? ')
    if has_another_skill == 'yes':
        skills = input('Add your skills ')
        p = d.add_paragraph(skills)
        p.style = 'List Bullet'
    else:
        break

d.save('cv.docx')
